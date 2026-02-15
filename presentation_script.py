"""
Presentation Script Generator
Converts the Jupyter notebook presentation content into a Word (.docx) document
with section-by-section speaking notes for a presentation on
"Predicting Student Course Completion Using Machine Learning".
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_section(doc, title, content, speaking_notes):
    """Add a presentation section with content and speaking notes."""
    # Section title
    heading = doc.add_heading(title, level=1)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)

    # Section content
    for paragraph_text in content:
        p = doc.add_paragraph(paragraph_text)
        for run in p.runs:
            run.font.size = Pt(11)

    # Speaking notes box
    doc.add_paragraph()  # spacer
    notes_heading = doc.add_paragraph()
    run = notes_heading.add_run("WHAT TO SAY:")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 102, 51)

    for note in speaking_notes:
        p = doc.add_paragraph(note, style="List Bullet")
        for run in p.runs:
            run.italic = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(80, 80, 80)

    # Section divider
    doc.add_paragraph("─" * 60)


def create_presentation_docx():
    """Create the presentation Word document with all sections and speaking notes."""
    doc = Document()

    # --- Title Page ---
    title = doc.add_heading("Predicting Student Course Completion\nUsing Machine Learning", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Clifton Chen Yi — 231220B")
    run.font.size = Pt(16)
    run.bold = True

    doc.add_paragraph()
    doc.add_paragraph("─" * 60)

    # =========================================================================
    # SECTION 1: Introduction / Overview
    # =========================================================================
    add_section(
        doc,
        title="1. Introduction & Problem Statement",
        content=[
            "Problem: Predict whether a student will complete an online course — a binary classification task (Completed vs Not Completed).",
            "Motivation: Online learning platforms face high dropout rates — often exceeding 90% on MOOCs. Early identification of at-risk students enables targeted interventions such as personalised reminders and additional support, which can significantly improve completion rates and platform revenue.",
            "Dataset: Student Course Completion Prediction Dataset from Kaggle — 100,000 student-course enrolment records with 40 features covering demographics, course metadata, engagement behaviour, and payment details.",
            "Success Criteria: F1-Score ≥ 0.70, ROC-AUC ≥ 0.75, Accuracy ≥ 0.70, and cross-validation standard deviation < 0.02.",
            "Approach: Train and compare three supervised classifiers — Logistic Regression, Random Forest, and Gradient Boosting — then tune the best performer and validate with Stratified K-Fold Cross-Validation.",
        ],
        speaking_notes=[
            "Good [morning/afternoon], today I'll be presenting my machine learning project on predicting student course completion.",
            "The problem I'm tackling is a real and important one: online courses have extremely high dropout rates, sometimes over 90%. If we can predict which students are likely to drop out early, platforms can intervene — send reminders, offer support — and improve outcomes.",
            "I used a Kaggle dataset with 100,000 records and 40 features. My goal was to achieve an F1-score of at least 0.70.",
            "I'll walk you through my entire pipeline: from data exploration to model training, comparison, tuning, and validation.",
        ],
    )

    # =========================================================================
    # SECTION 2: Data Exploration (EDA)
    # =========================================================================
    add_section(
        doc,
        title="2. Data Exploration (EDA)",
        content=[
            "The dataset contains 100,000 records and 40 columns with zero missing values — it is pre-cleaned.",
            "Target Variable: Nearly balanced — approximately 49% Completed vs 51% Not Completed. This means accuracy is a valid metric and class-imbalance techniques (e.g., SMOTE) are unnecessary.",
            "Numerical Features: Age is roughly uniform (17–40). Login frequency is right-skewed. Video completion rate varies widely (0–100%).",
            "Correlation Analysis: Most features show low inter-correlation, meaning features contribute largely independent information and multicollinearity is not a major concern.",
            "Boxplots by Completion Status: Progress Percentage shows the clearest separation between completers and non-completers. Video Completion Rate and Quiz Score Average also show meaningful differences.",
            "Categorical Features: Gender, Employment Status, and Device Type show remarkably similar completion rates — demographic features have limited predictive power.",
            "Outlier Analysis: Checked for outliers using the IQR method on key numerical features. Some features like Time Spent Hours and Payment Amount show outliers but they were retained as they represent genuine variation.",
        ],
        speaking_notes=[
            "Starting with exploratory data analysis — the dataset is quite large at 100,000 rows and comes pre-cleaned with no missing values.",
            "An important finding is that the target variable is nearly balanced at roughly 49/51, so I didn't need to worry about class imbalance techniques like SMOTE.",
            "The most interesting EDA insight is that engagement features — Progress Percentage, Video Completion Rate, Quiz Scores — show clear separation between completers and non-completers, while demographic features like gender and device type show almost no difference.",
            "This told me early on that behavioural engagement would be the key predictor, not who the student is, but how they interact with the course.",
        ],
    )

    # =========================================================================
    # SECTION 3: Data Cleaning & Preparation
    # =========================================================================
    add_section(
        doc,
        title="3. Data Cleaning & Preparation",
        content=[
            "Step 1 — Introduce Dirty Data: Since the dataset was pre-cleaned, I intentionally introduced approximately 2% missing values, 500 duplicate rows, and inconsistent capitalisations to practise real-world preprocessing.",
            "Step 2 — Clean the Data: Removed duplicates, imputed missing numerical values with median (chosen over mean due to skewed distributions), and standardised text formatting.",
            "Step 3 — Drop Non-Predictive Columns: Removed Student_ID, Name, Enrollment_Date, City, Course_ID, and Course_Name — these are identifiers or high-cardinality features with no predictive value.",
            "Step 4 — Target Encoding: Mapped 'Completed' → 1 and 'Not Completed' → 0.",
            "Step 5 — Categorical Encoding: Used a hybrid approach — ordinal encoding for naturally ordered features (Education Level, Course Level, Internet Connection Quality), one-hot encoding for low-cardinality nominal features (Gender, Employment Status, Device Type).",
            "Step 6 — Feature Engineering: Created 'Assignment_Completion_Rate' (ratio of submitted to total assignments) and 'Quiz_Performance' (interaction between quiz scores and quiz participation).",
            "Step 7 — Train-Test Split: 80/20 split with stratification on the target variable.",
            "Step 8 — Feature Scaling: Applied StandardScaler (mean=0, std=1) — essential for Logistic Regression; tree-based models are scale-invariant but scaling doesn't hurt them.",
        ],
        speaking_notes=[
            "For data cleaning, because the dataset came pre-cleaned, I deliberately introduced dirty data — missing values, duplicates, and inconsistencies — so I could demonstrate real-world preprocessing skills.",
            "A key decision was using median imputation instead of mean imputation, because several features like Login Frequency are right-skewed, and the median is more robust to outliers.",
            "I used a hybrid encoding strategy: ordinal encoding for features with a natural order like Education Level, and one-hot encoding for nominal features like Gender. This avoids the dimensionality explosion you'd get from one-hot encoding everything.",
            "I also engineered two new features: Assignment Completion Rate, which captures the ratio rather than raw counts, and Quiz Performance, which combines quiz scores with participation rate. These capture engagement quality beyond raw numbers.",
        ],
    )

    # =========================================================================
    # SECTION 4: Model Training
    # =========================================================================
    add_section(
        doc,
        title="4. Model Training",
        content=[
            "Three classification models were trained, representing different algorithm families:",
            "Model 1 — Logistic Regression: A linear model serving as an interpretable baseline. Training accuracy: 0.6070, Test accuracy: 0.6047. Minimal overfitting — the two values are very close.",
            "Model 2 — Random Forest: An ensemble of decision trees. Training accuracy: 1.0000, Test accuracy: 0.5933. Large train-test gap indicates significant overfitting — the model memorises training data but fails to generalise.",
            "Model 3 — Gradient Boosting: A sequential boosting ensemble. Training accuracy: 0.6307, Test accuracy: 0.6040. Moderate overfitting gap, but much less than Random Forest.",
        ],
        speaking_notes=[
            "I trained three different types of models to cover different algorithm families.",
            "Logistic Regression served as my baseline — it's simple, interpretable, and achieved about 60% accuracy with virtually no overfitting.",
            "Random Forest was interesting because it achieved perfect training accuracy of 100%, but only 59% on the test set — a clear sign of overfitting. The model was memorising the training data instead of learning generalisable patterns.",
            "Gradient Boosting fell in between, with a slight overfitting gap but better generalisation than Random Forest.",
            "An important observation here is that all three models achieved roughly 60% accuracy, which suggests the features have limited predictive signal — this is a dataset constraint, not a modelling failure.",
        ],
    )

    # =========================================================================
    # SECTION 5: Model Comparison
    # =========================================================================
    add_section(
        doc,
        title="5. Model Comparison",
        content=[
            "All models were compared using: Accuracy, Precision, Recall, F1-Score, and ROC-AUC.",
            "Results Summary:",
            "• Logistic Regression — Accuracy: 0.6047, F1: 0.5924, AUC: 0.6484 (Best overall)",
            "• Random Forest — Accuracy: 0.5933, F1: 0.5880, AUC: 0.6316",
            "• Gradient Boosting — Accuracy: 0.6040, F1: 0.5932, AUC: 0.6488",
            "All models achieved modest performance (F1 ≈ 0.58–0.59, AUC ≈ 0.63–0.65), falling short of the initial success criteria (F1 ≥ 0.70, AUC ≥ 0.75).",
            "Confusion Matrix Analysis: All models show a fairly even distribution of errors between false positives and false negatives, consistent with the balanced dataset.",
            "ROC Curve Analysis: AUC values around 0.63–0.65 indicate the models perform better than random (0.5) but have limited discriminative ability.",
            "Feature Importance: Behavioural engagement features (Progress Percentage, Video Completion Rate, Quiz Score Average) dominate — confirming EDA findings.",
        ],
        speaking_notes=[
            "Comparing all three models across multiple metrics, the results are quite similar — all hover around 60% accuracy and an F1-score of about 0.59.",
            "Logistic Regression slightly outperformed the tree-based models, which is notable because it's the simplest model. This suggests the relationships in this data are largely linear.",
            "None of the models met my initial success criteria of F1 ≥ 0.70, which points to a fundamental limitation in the dataset's features rather than a modelling issue.",
            "Feature importance from Random Forest confirmed what we saw in EDA — engagement behaviours like progress percentage and video completion rate are far more predictive than demographics.",
            "I selected Gradient Boosting for tuning because it has the most tunable hyperparameters and achieved competitive performance.",
        ],
    )

    # =========================================================================
    # SECTION 6: Hyperparameter Tuning
    # =========================================================================
    add_section(
        doc,
        title="6. Hyperparameter Tuning",
        content=[
            "Method: GridSearchCV with 3-fold Stratified Cross-Validation on Gradient Boosting.",
            "Hyperparameters tuned: n_estimators (50, 100), max_depth (3, 5, 7), learning_rate (0.01, 0.1).",
            "Best parameters found: learning_rate=0.1, max_depth=3, n_estimators=100.",
            "Key insight: A shallower tree depth (3 vs the original 5) was preferred, confirming that reducing model complexity improves generalisation.",
            "Result: Tuning selected shallower trees but did not meaningfully improve test performance, confirming the default parameters were near-optimal.",
        ],
        speaking_notes=[
            "For hyperparameter tuning, I used GridSearchCV with 3-fold cross-validation on the Gradient Boosting model.",
            "The most interesting finding was that the grid search preferred shallower trees — max depth of 3 instead of 5. This confirms that simpler models generalise better on this dataset.",
            "However, tuning didn't significantly improve performance, which reinforces the conclusion that the performance ceiling is due to the dataset's limited predictive signal, not suboptimal hyperparameters.",
        ],
    )

    # =========================================================================
    # SECTION 7: Validation
    # =========================================================================
    add_section(
        doc,
        title="7. Validation (Cross-Validation)",
        content=[
            "Method: 5-Fold Stratified Cross-Validation on the tuned Gradient Boosting model.",
            "Results: Accuracy ≈ 0.60, F1 ≈ 0.59, Precision ≈ 0.59, Recall ≈ 0.59.",
            "All metrics show low standard deviation (< 0.02), confirming stable generalisation across data splits.",
            "Consistent performance across folds demonstrates the model is not overfitting to any particular data split.",
        ],
        speaking_notes=[
            "To validate the model's robustness, I used 5-fold stratified cross-validation.",
            "The key takeaway is that all metrics — accuracy, F1, precision, recall — show very low standard deviation, less than 0.02. This means the model performs consistently regardless of how the data is split.",
            "This meets my cross-validation stability criterion and confirms the model generalises well, even though the absolute performance is modest.",
        ],
    )

    # =========================================================================
    # SECTION 8: Conclusion & Recommendations
    # =========================================================================
    add_section(
        doc,
        title="8. Conclusion & Recommendations",
        content=[
            "Summary: Built a complete binary classification pipeline with 100,000 records and 40 features. All models achieved ~60% accuracy (F1 ≈ 0.59, AUC ≈ 0.65).",
            "Key Decision Points:",
            "• Hybrid encoding strategy (ordinal + one-hot) to avoid dimensionality explosion.",
            "• Chose Logistic Regression, Random Forest, and Gradient Boosting over SVM due to the 100K-row dataset making SVM computationally expensive.",
            "Limitations:",
            "• Limited predictive signal from the 40 features — all models plateau at ~60% accuracy.",
            "• Potential data leakage from Progress_Percentage (may partially encode completion status).",
            "• Dataset appears synthetic — real-world data would contain more noise and imbalance.",
            "Recommendations for Course Providers:",
            "1. Use engagement metrics (Progress Percentage, Video Completion Rate, Quiz Scores) to identify at-risk students early.",
            "2. Trigger automated interventions when predicted completion probability drops below a threshold.",
            "3. Improve course design by analysing which engagement factors most influence completion.",
            "4. Collect additional data (forum participation, video watch patterns, prior course history) to improve prediction accuracy beyond 60%.",
        ],
        speaking_notes=[
            "To conclude, I built a complete, rigorous machine learning pipeline from data exploration through to validated predictions.",
            "While the models didn't meet my initial success criteria of F1 ≥ 0.70, this is an important finding in itself — it tells us that the available 40 features have limited power to predict course completion.",
            "The practical takeaway is that engagement metrics — especially Progress Percentage, Video Completion Rate, and Quiz Scores — are the strongest predictors. Course providers should monitor these in real time and trigger interventions when they drop.",
            "For future work, I'd recommend collecting richer behavioural data — like forum participation and detailed video watch patterns — which could push prediction accuracy well beyond the current 60%.",
            "Thank you for listening. I'm happy to take any questions.",
        ],
    )

    # Save the document
    output_path = "Presentation_Script.docx"
    doc.save(output_path)
    print(f"Presentation script saved to: {output_path}")
    return output_path


if __name__ == "__main__":
    create_presentation_docx()
